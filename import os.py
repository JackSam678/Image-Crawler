# import os
# import shutil
# import requests
# from PIL import Image
# from docx import Document
# from docx.shared import Inches
# from openpyxl import Workbook
# from openpyxl.drawing.image import Image as ExcelImage
# from selenium import webdriver
# from selenium.webdriver.common.by import By
# from selenium.webdriver.firefox.service import Service
# import time

# # 配置参数（请确保URL正确）
# URL = "https://www.yaash.cn/desk-wallpaper/?orderby=views"  # 示例：替换为你需要爬取的网页
# SAVE_DIR = os.path.join(os.path.expanduser("~"), "桌面", "PC_Test")
# ORIGINAL_IMG_DIR = os.path.join(SAVE_DIR, "original_images")
# CONVERTED_IMG_DIR = os.path.join(SAVE_DIR, "converted_images")
# MIN_IMG_SIZE = 30000
# LOAD_TIMES = 5
# GECKODRIVER_PATH = "../geckodriver-v0.36.0-linux64/geckodriver"

# # 初始化目录
# def init_dirs():
#     for dir_path in [SAVE_DIR, ORIGINAL_IMG_DIR, CONVERTED_IMG_DIR]:
#         os.makedirs(dir_path, exist_ok=True)
#     print(f"工作目录初始化完成: {SAVE_DIR}")

# # 滑动加载页面
# def load_page(driver):
#     try:
#         driver.get(URL)
#         print(f"已成功访问网页：{URL}")
#         time.sleep(3)  # 延长初始加载时间
#     except Exception as e:
#         print(f"访问网页失败：{str(e)}")
#         return driver
    
#     for i in range(LOAD_TIMES):
#         try:
#             driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
#             print(f"滑动加载第{i+1}次")
#             time.sleep(3)  # 延长滑动后等待时间
#         except Exception as e:
#             print(f"滑动加载失败：{str(e)}")
#             break
#     return driver

# # 提取文字段落（适配网页结构）
# def extract_text(driver):
#     try:
#         # 尝试多种标签提取文字（根据网页调整）
#         selectors = [By.TAG_NAME, "p", By.TAG_NAME, "div", By.TAG_NAME, "span"]
#         text_list = []
#         for i in range(0, len(selectors), 2):
#             by = selectors[i]
#             tag = selectors[i+1]
#             elements = driver.find_elements(by, tag)
#             text_list.extend([e.text.strip() for e in elements if e.text.strip()])
#         # 去重
#         text_list = list(set(text_list))
#         print(f"提取到{len(text_list)}段文字")
#         return text_list
#     except Exception as e:
#         print(f"提取文字失败：{str(e)}")
#         return []

# # 提取图片链接
# def extract_images(driver):
#     try:
#         img_tags = driver.find_elements(By.TAG_NAME, "img")
#         img_urls = []
#         for img in img_tags:
#             url = img.get_attribute("src")
#             if url:
#                 # 补全相对路径URL
#                 if url.startswith("//"):
#                     url = "https:" + url
#                 elif url.startswith("/"):
#                     url = f"{URL}{url}"
#                 if "avif" not in url.lower():
#                     img_urls.append(url)
#         # 去重
#         img_urls = list(set(img_urls))
#         print(f"提取到{len(img_urls)}张图片链接（已过滤AVIF）")
#         return img_urls
#     except Exception as e:
#         print(f"提取图片链接失败：{str(e)}")
#         return []

# # 下载图片
# def download_images(img_urls):
#     saved_paths = []
#     for i, url in enumerate(img_urls, 1):
#         try:
#             print(f"正在下载图片{i}/{len(img_urls)}：{url}")
#             headers = {
#                 "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:129.0) Gecko/20100101 Firefox/129.0",
#                 "Referer": URL  # 添加来源页，避免反爬
#             }
#             response = requests.get(url, headers=headers, timeout=15)
#             response.raise_for_status()
            
#             if len(response.content) < MIN_IMG_SIZE:
#                 print(f"图片过小（{len(response.content)}字节），跳过")
#                 continue
            
#             ext = os.path.splitext(url.split('?')[0])[-1] or ".jpg"
#             img_name = f"img_{i}{ext}"
#             original_path = os.path.join(ORIGINAL_IMG_DIR, img_name)
#             with open(original_path, "wb") as f:
#                 f.write(response.content)
            
#             saved_paths.append(original_path)
#             print(f"图片{i}保存成功")
#         except Exception as e:
#             print(f"图片{i}下载失败：{str(e)}")
#     print(f"成功下载{len(saved_paths)}/{len(img_urls)}张图片")
#     return saved_paths

# # 转换图片
# def convert_images(original_paths):
#     converted_paths = []
#     for i, original_path in enumerate(original_paths, 1):
#         try:
#             with Image.open(original_path) as img:
#                 if img.mode in ("RGBA", "P"):
#                     img = img.convert("RGB")
#                 converted_path = os.path.join(CONVERTED_IMG_DIR, f"converted_{i}.jpg")
#                 img.save(converted_path, "JPEG")
#                 converted_paths.append(converted_path)
#                 print(f"图片{i}转换为JPG成功：{converted_path}")
#         except Exception as e:
#             converted_path = os.path.join(CONVERTED_IMG_DIR, f"converted_{i}{os.path.splitext(original_path)[-1]}")
#             shutil.copy2(original_path, converted_path)
#             converted_paths.append(converted_path)
#             print(f"图片{i}转换失败：{str(e)}，已复制原始文件备用")
#     return converted_paths

# # 保存为Word
# def save_to_word(texts, images):
#     try:
#         doc = Document()
#         for i, text in enumerate(texts, 1):
#             doc.add_paragraph(text)
#             print(f"添加文字段落{i}")
#         for i, img_path in enumerate(images, 1):
#             try:
#                 doc.add_picture(img_path, width=Inches(5))
#                 print(f"添加图片{i}到文档")
#             except Exception as e:
#                 print(f"图片{i}插入失败：{str(e)}")
#         doc_path = os.path.join(SAVE_DIR, "网页内容汇总.docx")
#         doc.save(doc_path)
#         print(f"文档保存成功：{doc_path}")
#     except Exception as e:
#         print(f"保存Word失败：{str(e)}")

# # 保存为Excel
# def save_to_excel(texts, images):
#     try:
#         wb = Workbook()
#         ws = wb.active
#         ws.title = "网页内容汇总"
#         ws.append(["文字内容"])
#         for i, text in enumerate(texts, 1):
#             ws.cell(row=i+1, column=1, value=text)
#             ws.row_dimensions[i+1].height = 40
#             print(f"Excel添加文字段落{i}")
#         current_row = len(texts) + 3
#         ws.cell(row=current_row-1, column=1, value="图片内容")
#         for i, img_path in enumerate(images, 1):
#             try:
#                 img = ExcelImage(img_path)
#                 img.width = 400
#                 img.height = 300
#                 ws.add_image(img, f"A{current_row}")
#                 ws.row_dimensions[current_row].height = 250
#                 current_row += 10
#                 print(f"Excel添加图片{i}")
#             except Exception as e:
#                 print(f"Excel图片{i}插入失败：{str(e)}")
#         excel_path = os.path.join(SAVE_DIR, "网页内容汇总.xlsx")
#         wb.save(excel_path)
#         print(f"Excel保存成功：{excel_path}")
#     except Exception as e:
#         print(f"保存Excel失败：{str(e)}")

# # 主函数
# def main():
#     init_dirs()
    
#     if not os.path.exists(GECKODRIVER_PATH):
#         print(f"错误：驱动文件不存在 - {GECKODRIVER_PATH}")
#         return
    
#     try:
#         # 简化驱动配置，使用已验证的正常驱动
#         service = Service(executable_path=GECKODRIVER_PATH)
#         driver = webdriver.Firefox(service=service)
#         # 最大化窗口，确保内容加载完整
#         driver.maximize_window()
#         # 设置超时时间
#         driver.set_page_load_timeout(30)
#         driver.implicitly_wait(10)  # 元素查找超时
#     except Exception as e:
#         print(f"启动火狐浏览器失败：{str(e)}")
#         return
    
#     try:
#         driver = load_page(driver)
#         texts = extract_text(driver)
#         img_urls = extract_images(driver)
#         original_img_paths = download_images(img_urls)
#         converted_img_paths = convert_images(original_img_paths)
#         save_to_word(texts, converted_img_paths)
#         save_to_excel(texts, converted_img_paths)
#         print(f"原始图片目录：{ORIGINAL_IMG_DIR}（共{len(os.listdir(ORIGINAL_IMG_DIR))}个文件）")
#         print(f"转换后图片目录：{CONVERTED_IMG_DIR}（共{len(os.listdir(CONVERTED_IMG_DIR))}个文件）")
#     finally:
#         # 确保浏览器关闭
#         try:
#             driver.quit()
#         except:
#             pass

# if __name__ == "__main__":
#     main()

###################################################################################################################

#优化后的代码
import os
import shutil
import requests
import time
from datetime import datetime
from PIL import Image
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from selenium.webdriver.edge.service import Service as EdgeService


# -------------------------- 配置参数 --------------------------
# 浏览器选择："chrome" / "firefox" / "edge"（自动检测可用驱动）
BROWSER = "edge"  # 默认用火狐

# 驱动路径配置（请根据实际路径修改）
DRIVER_PATHS = {
    "chrome": "../chromedriver",  # 谷歌驱动路径
    "firefox": "../geckodriver",  # 火狐驱动路径
    "edge": "../msedgedriver"  # Edge驱动路径
}

# 爬取目标URL（每次运行可修改）
URL = "https://www.yaash.cn/desk-wallpaper/page/4/"

# 其他配置
MIN_IMG_SIZE = 30000  # 最小图片字节数
LOAD_TIMES = 5  # 滑动加载次数
BASE_SAVE_DIR = os.path.join(os.path.expanduser("~"), "桌面", "PC_Test")  # 基础保存目录
# --------------------------------------------------------------


def get_timestamp_dir():
    """生成带时间戳的目录名，确保每次运行保存到新目录"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(BASE_SAVE_DIR, f"crawl_{timestamp}")


def init_dirs(save_dir):
    """初始化目录（带时间戳的新目录）"""
    dirs = {
        "original": os.path.join(save_dir, "original_images"),
        "converted": os.path.join(save_dir, "converted_images")
    }
    for dir_path in dirs.values():
        os.makedirs(dir_path, exist_ok=True)
    print(f"工作目录初始化完成: {save_dir}")
    return dirs


def get_browser_driver(browser):
    """根据浏览器类型初始化驱动（自动检测可用驱动）"""
    driver_path = DRIVER_PATHS.get(browser)
    
    # 检查驱动是否存在
    if not os.path.exists(driver_path):
        print(f"警告：{browser}驱动路径不存在 - {driver_path}，尝试自动查找系统驱动...")
        driver_path = None  # 让Selenium自动查找系统PATH中的驱动
    
    try:
        if browser == "chrome":
            service = ChromeService(executable_path=driver_path) if driver_path else ChromeService()
            options = webdriver.ChromeOptions()
            options.add_argument("--disable-blink-features=AutomationControlled")
            return webdriver.Chrome(service=service, options=options)
        
        elif browser == "firefox":
            service = FirefoxService(executable_path=driver_path) if driver_path else FirefoxService()
            options = webdriver.FirefoxOptions()
            options.add_argument("--disable-blink-features=AutomationControlled")
            return webdriver.Firefox(service=service, options=options)
        
        elif browser == "edge":
            service = EdgeService(executable_path=driver_path) if driver_path else EdgeService()
            options = webdriver.EdgeOptions()
            options.add_argument("--disable-blink-features=AutomationControlled")
            return webdriver.Edge(service=service, options=options)
        
        else:
            raise ValueError(f"不支持的浏览器：{browser}，可选：chrome/firefox/edge")
    
    except Exception as e:
        print(f"{browser}驱动初始化失败：{str(e)}")
        return None


def load_page(driver, url):
    """滑动加载页面内容"""
    try:
        driver.get(url)
        print(f"已成功访问网页：{url}")
        time.sleep(3)  # 初始加载等待
    except Exception as e:
        print(f"访问网页失败：{str(e)}")
        return driver
    
    # 滑动加载更多内容
    for i in range(LOAD_TIMES):
        try:
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            print(f"滑动加载第{i+1}次")
            time.sleep(3)  # 等待加载
        except Exception as e:
            print(f"滑动加载失败：{str(e)}")
            break
    return driver


def extract_text(driver):
    """提取网页文字内容"""
    try:
        # 适配不同网页的文字标签
        selectors = [
            (By.TAG_NAME, "p"),
            (By.TAG_NAME, "div"),
            (By.TAG_NAME, "span"),
            (By.TAG_NAME, "h1"),
            (By.TAG_NAME, "h2")
        ]
        text_list = []
        for by, tag in selectors:
            elements = driver.find_elements(by, tag)
            text_list.extend([e.text.strip() for e in elements if e.text.strip()])
        
        # 去重并过滤过短文本
        text_list = list(filter(lambda x: len(x) > 5, list(set(text_list))))
        print(f"提取到{len(text_list)}段有效文字")
        return text_list
    except Exception as e:
        print(f"提取文字失败：{str(e)}")
        return []


def extract_images(driver, base_url):
    """提取图片链接（补全相对路径）"""
    try:
        img_tags = driver.find_elements(By.TAG_NAME, "img")
        img_urls = []
        for img in img_tags:
            url = img.get_attribute("src") or img.get_attribute("data-src")  # 处理懒加载图片
            if not url:
                continue
            
            # 补全相对路径
            if url.startswith("//"):
                url = "https:" + url
            elif url.startswith("/") and not url.startswith("//"):
                url = base_url.rstrip("/") + url
            elif not url.startswith(("http://", "https://")):
                continue  # 忽略无效链接
            
            # 过滤AVIF格式
            if "avif" not in url.lower():
                img_urls.append(url)
        
        # 去重
        img_urls = list(set(img_urls))
        print(f"提取到{len(img_urls)}张图片链接（已过滤AVIF）")
        return img_urls
    except Exception as e:
        print(f"提取图片链接失败：{str(e)}")
        return []


def download_images(img_urls, save_dir):
    """下载图片到原始图片目录"""
    saved_paths = []
    for i, url in enumerate(img_urls, 1):
        try:
            print(f"正在下载图片{i}/{len(img_urls)}：{url}")
            headers = {
                "User-Agent": "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:136.0) Gecko/20100101 Firefox/136.0",
                "Referer": URL
            }
            response = requests.get(url, headers=headers, timeout=20)
            response.raise_for_status()
            
            # 过滤小图片
            if len(response.content) < MIN_IMG_SIZE:
                print(f"图片过小（{len(response.content)}字节），跳过")
                continue
            
            # 保存图片（保留原始扩展名）
            ext = os.path.splitext(url.split('?')[0])[-1] or ".jpg"
            img_name = f"img_{i}{ext}"
            save_path = os.path.join(save_dir, img_name)
            with open(save_path, "wb") as f:
                f.write(response.content)
            
            saved_paths.append(save_path)
            print(f"图片{i}保存成功")
        except Exception as e:
            print(f"图片{i}下载失败：{str(e)}")
    
    print(f"成功下载{len(saved_paths)}/{len(img_urls)}张图片")
    return saved_paths


def convert_images(original_paths, converted_dir):
    """转换图片为JPG格式"""
    converted_paths = []
    for i, original_path in enumerate(original_paths, 1):
        try:
            with Image.open(original_path) as img:
                # 处理透明通道
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                
                converted_name = f"converted_{i}.jpg"
                converted_path = os.path.join(converted_dir, converted_name)
                img.save(converted_path, "JPEG")
                converted_paths.append(converted_path)
                print(f"图片{i}转换为JPG成功：{converted_path}")
        except Exception as e:
            # 转换失败时复制原始文件
            converted_name = f"converted_{i}{os.path.splitext(original_path)[-1]}"
            converted_path = os.path.join(converted_dir, converted_name)
            shutil.copy2(original_path, converted_path)
            converted_paths.append(converted_path)
            print(f"图片{i}转换失败：{str(e)}，已复制原始文件备用")
    return converted_paths


def save_to_word(texts, images, save_dir):
    """保存内容到Word文档（带时间戳）"""
    try:
        doc = Document()
        doc.add_heading("网页内容汇总", level=1)
        doc.add_paragraph(f"爬取时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"爬取URL：{URL}")
        doc.add_paragraph("--- 文字内容 ---")
        
        # 添加文字段落
        for i, text in enumerate(texts, 1):
            doc.add_paragraph(f"段落{i}：{text}")
            print(f"添加文字段落{i}")
        
        # 添加图片
        doc.add_page_break()
        doc.add_paragraph("--- 图片内容 ---")
        for i, img_path in enumerate(images, 1):
            try:
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph(f"图片{i}：{os.path.basename(img_path)}")
                print(f"添加图片{i}到文档")
            except Exception as e:
                print(f"图片{i}插入失败：{str(e)}")
        
        # 保存文档（带时间戳）
        doc_path = os.path.join(save_dir, "网页内容汇总.docx")
        doc.save(doc_path)
        print(f"Word文档保存成功：{doc_path}")
    except Exception as e:
        print(f"保存Word失败：{str(e)}")


def save_to_excel(texts, images, save_dir):
    """保存内容到Excel文档（带时间戳）"""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "网页内容汇总"
        
        # 基本信息
        ws.append(["爬取信息", ""])
        ws.append(["爬取时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        ws.append(["爬取URL", URL])
        ws.append(["--- 文字内容 ---", ""])
        
        # 文字段落
        for i, text in enumerate(texts, 1):
            ws.append([f"段落{i}", text])
            ws.row_dimensions[i+4].height = 40  # 调整行高
            print(f"Excel添加文字段落{i}")
        
        # 图片内容
        ws.append(["--- 图片内容 ---", ""])
        current_row = len(texts) + 6
        for i, img_path in enumerate(images, 1):
            try:
                ws.cell(row=current_row, column=1, value=f"图片{i}")
                img = ExcelImage(img_path)
                img.width = 400
                img.height = 300
                ws.add_image(img, f"B{current_row}")
                ws.row_dimensions[current_row].height = 250
                current_row += 15  # 间隔15行放一张图
                print(f"Excel添加图片{i}")
            except Exception as e:
                print(f"Excel图片{i}插入失败：{str(e)}")
        
        # 保存Excel（带时间戳）
        excel_path = os.path.join(save_dir, "网页内容汇总.xlsx")
        wb.save(excel_path)
        print(f"Excel文档保存成功：{excel_path}")
    except Exception as e:
        print(f"保存Excel失败：{str(e)}")


def main():
    # 生成带时间戳的新目录（每次运行不同）
    save_dir = get_timestamp_dir()
    dirs = init_dirs(save_dir)  # dirs = {"original":..., "converted":...}
    
    # 初始化浏览器驱动
    driver = get_browser_driver(BROWSER)
    if not driver:
        print("无法初始化浏览器驱动，程序退出")
        return
    
    try:
        driver.maximize_window()
        driver.set_page_load_timeout(60)  # 延长超时时间
        driver.implicitly_wait(10)
        
        # 爬取流程
        driver = load_page(driver, URL)
        texts = extract_text(driver)
        img_urls = extract_images(driver, URL)
        original_img_paths = download_images(img_urls, dirs["original"])
        converted_img_paths = convert_images(original_img_paths, dirs["converted"])
        
        # 保存文档（每次都是新文件）
        save_to_word(texts, converted_img_paths, save_dir)
        save_to_excel(texts, converted_img_paths, save_dir)
        
        # 输出结果统计
        print(f"\n爬取完成！所有文件保存至：{save_dir}")
        print(f"原始图片：{len(os.listdir(dirs['original']))}个")
        print(f"转换后图片：{len(os.listdir(dirs['converted']))}个")
    
    finally:
        driver.quit()


if __name__ == "__main__":
    main()
