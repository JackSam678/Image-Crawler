import os
import time
import re
import requests
import shutil
from datetime import datetime
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException


# -------------------------- 配置参数 --------------------------
# 目标网页URL（替换为任意网页）
TARGET_URL = ""

# 浏览器配置（chrome/firefox/edge）
BROWSER = "edge"

# 驱动路径（根据实际存放位置修改）
DRIVER_PATHS = {
    "chrome": "/Image Crawler/Drive location/chromedriver-linux64/chromedriver",
    "firefox": "/Image Crawler/Drive location/geckodriver-v0.36.0-linux64/geckodriver",
    "edge": "/Image Crawler/Drive location/edgedriver_linux64/msedgedriver"
}

# 爬取配置
SCROLL_TIMES = 3  # 滚动加载次数（动态内容较多可增加）
SCROLL_INTERVAL = 2  # 滚动间隔（秒）
MIN_IMAGE_SIZE = 10000  # 最小图片大小（字节，降低阈值避免漏爬）
TIMEOUT = 15  # 元素等待超时时间（秒）
MIN_TEXT_LENGTH = 20  # 最小文本长度（过滤无意义短句）

# 保存路径：当前脚本所在文件夹
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
# --------------------------------------------------------------


class UniversalWebCrawler:
    def __init__(self):
        # 初始化保存目录（网页爬取数据_时间戳）
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.save_dir = os.path.join(CURRENT_DIR, f"网页爬取数据_{self.timestamp}")
        self.image_dir = os.path.join(self.save_dir, "images")
        self.init_dirs()
        
        # 解析目标网站域名（用于补全相对路径）
        self.base_domain = f"{urlparse(TARGET_URL).scheme}://{urlparse(TARGET_URL).netloc}"
        
        # 初始化浏览器驱动
        self.driver = self.init_browser()
        if not self.driver:
            raise Exception("浏览器驱动初始化失败")

    def init_dirs(self):
        """创建保存目录"""
        for dir_path in [self.save_dir, self.image_dir]:
            os.makedirs(dir_path, exist_ok=True)
        print(f"数据保存目录：{self.save_dir}")

    def init_browser(self):
        """初始化浏览器（模拟真实用户）"""
        try:
            driver_path = DRIVER_PATHS.get(BROWSER)
            if driver_path and not os.path.exists(driver_path):
                print(f"警告：未找到驱动文件，尝试自动查找系统驱动...")
                driver_path = None

            # 通用配置（适配大多数网站反爬）
            if BROWSER == "chrome":
                options = webdriver.ChromeOptions()
                # 关闭无头模式（调试时用，稳定后可开启）
                # options.add_argument("--headless=new")
                options.add_argument("--disable-blink-features=AutomationControlled")
                options.add_argument("--disable-gpu")
                options.add_argument("--no-sandbox")
                options.add_argument("window-size=1920,1080")
                options.add_experimental_option("excludeSwitches", ["enable-automation"])
                options.add_experimental_option("useAutomationExtension", False)
                service = ChromeService(executable_path=driver_path) if driver_path else ChromeService()
                driver = webdriver.Chrome(service=service, options=options)
                # 去除 navigator.webdriver 标记
                driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
                    "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
                })
                return driver

            elif BROWSER == "firefox":
                options = webdriver.FirefoxOptions()
                # options.add_argument("--headless")
                options.add_argument("--disable-blink-features=AutomationControlled")
                options.add_argument("window-size=1920,1080")
                service = FirefoxService(executable_path=driver_path) if driver_path else FirefoxService()
                return webdriver.Firefox(service=service, options=options)

            elif BROWSER == "edge":
                options = webdriver.EdgeOptions()
                # options.add_argument("--headless=new")
                options.add_argument("--disable-blink-features=AutomationControlled")
                options.add_argument("window-size=1920,1080")
                service = EdgeService(executable_path=driver_path) if driver_path else EdgeService()
                driver = webdriver.Edge(service=service, options=options)
                driver.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument", {
                    "source": "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
                })
                return driver

            else:
                raise ValueError(f"不支持的浏览器：{BROWSER}")

        except Exception as e:
            print(f"驱动初始化失败：{str(e)}")
            return None

    def load_page(self):
        """智能加载页面（处理动态内容）"""
        try:
            print(f"访问网页：{TARGET_URL}")
            self.driver.get(TARGET_URL)
            time.sleep(3)  # 初始加载等待

            # 滚动加载（逐步滚动，避免一次性到底导致内容未加载）
            last_height = self.driver.execute_script("return document.body.scrollHeight")
            for i in range(SCROLL_TIMES):
                # 滚动到当前页面底部
                self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                print(f"滚动加载第{i+1}/{SCROLL_TIMES}次")
                time.sleep(SCROLL_INTERVAL)
                
                # 检查是否加载了新内容
                new_height = self.driver.execute_script("return document.body.scrollHeight")
                if new_height == last_height:
                    print("已滚动到页面底部，停止滚动")
                    break
                last_height = new_height

            # 等待主体内容加载（自适应标签：优先id为content/main的区域）
            content_tags = ["content", "main", "article", "container", "wrapper"]
            for tag in content_tags:
                try:
                    WebDriverWait(self.driver, TIMEOUT).until(
                        EC.presence_of_element_located((By.ID, tag))
                    )
                    print(f"检测到主体内容区域（id={tag}），加载完成")
                    return True
                except TimeoutException:
                    continue
            
            # 若未找到特定id，等待body加载完成
            WebDriverWait(self.driver, TIMEOUT).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            print("页面加载完成（未检测到特定内容区域）")
            return True

        except Exception as e:
            print(f"页面加载失败：{str(e)}")
            return False

    def extract_text(self):
        """智能提取文本（适配大多数网页结构）"""
        try:
            text_elements = []
            
            # 策略1：提取主体内容区域内的文本（优先id为content/main等的区域）
            content_tags = ["content", "main", "article", "container", "wrapper"]
            for tag in content_tags:
                content = self.driver.find_elements(By.ID, tag)
                if content:
                    text_elements.extend(content[0].find_elements(By.XPATH, ".//*[self::p or self::h1 or self::h2 or self::h3 or self::div]"))
                    break  # 找到一个有效区域即可
            
            # 策略2：若未找到主体区域，提取全页面关键标签文本
            if not text_elements:
                text_elements = self.driver.find_elements(
                    By.XPATH, "//p | //h1 | //h2 | //h3 | //h4 | //div[not(contains(@class, 'ad')) and not(contains(@class, 'footer'))]"
                )

            # 过滤和清洗文本
            text_list = []
            for elem in text_elements:
                text = elem.text.strip()
                # 过滤条件：长度达标 + 不是纯数字/特殊字符
                if text and len(text) >= MIN_TEXT_LENGTH and re.search(r'[^\d\s.,!?;:"\'<>()\[\]{}]', text):
                    text_list.append(text)
            
            # 去重（保留顺序）
            seen = set()
            unique_texts = []
            for text in text_list:
                if text not in seen:
                    seen.add(text)
                    unique_texts.append(text)
            
            print(f"提取到 {len(unique_texts)} 条有效文本")
            return unique_texts

        except Exception as e:
            print(f"文本提取失败：{str(e)}")
            return []

    def extract_images(self):
        """智能提取图片链接（补全路径+过滤无效图）"""
        try:
            img_urls = []
            # 提取所有img标签，排除明显的图标（通过尺寸或class判断）
            img_tags = self.driver.find_elements(
                By.XPATH, "//img[not(contains(@class, 'icon')) and not(contains(@class, 'logo')) and not(contains(@src, 'icon'))]"
            )
            
            for img in img_tags:
                # 尝试获取各种图片链接属性
                url = img.get_attribute("src") or img.get_attribute("data-src") or img.get_attribute("data-original")
                if not url:
                    continue
                
                # 智能补全路径
                if url.startswith("//"):
                    url = f"https:{url}"
                elif url.startswith("/"):
                    url = urljoin(self.base_domain, url)
                elif not url.startswith(("http://", "https://")):
                    url = urljoin(TARGET_URL, url)  # 相对路径拼接
                
                # 过滤过小的图片（通过尺寸属性预判）
                width = img.get_attribute("width") or 0
                height = img.get_attribute("height") or 0
                if width.isdigit() and height.isdigit() and (int(width) < 50 or int(height) < 50):
                    continue  # 过滤小图标
                
                img_urls.append(url)
            
            # 去重
            img_urls = list(set(img_urls))
            print(f"提取到 {len(img_urls)} 个有效图片链接")
            return img_urls

        except Exception as e:
            print(f"图片提取失败：{str(e)}")
            return []

    def download_images(self, img_urls):
        """下载图片（带重试机制）"""
        downloaded = []
        for i, url in enumerate(img_urls, 1):
            for retry in range(2):  # 最多重试2次
                try:
                    print(f"下载图片 {i}/{len(img_urls)}（重试{retry}次）：{url}")
                    headers = {
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36",
                        "Referer": self.base_domain,
                        "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8"
                    }
                    
                    response = requests.get(url, headers=headers, timeout=15, stream=True)
                    response.raise_for_status()
                    
                    # 检查文件大小（最终过滤）
                    content_length = int(response.headers.get("Content-Length", 0))
                    if content_length < MIN_IMAGE_SIZE:
                        print(f"图片过小（{content_length}字节），跳过")
                        break
                    
                    # 保存图片
                    ext = os.path.splitext(url.split("?")[0])[-1].lower()
                    if not ext or ext not in [".jpg", ".jpeg", ".png", ".gif", ".webp"]:
                        ext = ".jpg"  # 默认格式
                    img_name = f"image_{i}{ext}"
                    save_path = os.path.join(self.image_dir, img_name)
                    
                    with open(save_path, "wb") as f:
                        for chunk in response.iter_content(chunk_size=1024):
                            if chunk:
                                f.write(chunk)
                    
                    downloaded.append(save_path)
                    print(f"图片 {i} 下载成功")
                    break  # 成功则退出重试循环
                
                except Exception as e:
                    if retry < 1:
                        print(f"下载失败，重试...（错误：{str(e)}）")
                        time.sleep(1)
                    else:
                        print(f"图片 {i} 下载失败（已达最大重试次数）：{str(e)}")
        
        print(f"图片下载完成，成功 {len(downloaded)}/{len(img_urls)} 张")
        return downloaded

    def save_to_word(self, texts, images):
        """保存为Word文档"""
        try:
            doc = Document()
            doc.add_heading("网页爬取数据汇总", level=1)
            doc.add_paragraph(f"爬取时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"爬取URL：{TARGET_URL}")
            
            # 添加文本
            if texts:
                doc.add_heading("\n文本内容", level=2)
                for i, text in enumerate(texts, 1):
                    doc.add_paragraph(f"【段落{i}】{text}")
                    doc.add_paragraph("—"*50)  # 分隔线
            
            # 添加图片
            if images:
                doc.add_page_break()
                doc.add_heading("图片内容", level=2)
                for i, img_path in enumerate(images, 1):
                    try:
                        doc.add_picture(img_path, width=Inches(5))
                        doc.add_paragraph(f"【图片{i}】{os.path.basename(img_path)}")
                    except Exception as e:
                        print(f"图片 {i} 插入Word失败：{str(e)}")
            
            doc_path = os.path.join(self.save_dir, "爬取数据汇总.docx")
            doc.save(doc_path)
            print(f"Word文档已保存：{doc_path}")
        except Exception as e:
            print(f"保存Word失败：{str(e)}")

    def save_to_excel(self, texts, images):
        """保存为Excel表格"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "爬取数据"
            
            # 基础信息
            ws.append(["项目", "详情"])
            ws.append(["爬取时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
            ws.append(["爬取URL", TARGET_URL])
            ws.append(["文本数量", len(texts)])
            ws.append(["图片数量", len(images)])
            ws.append([])
            
            # 文本内容
            if texts:
                ws.append(["文本内容"])
                for i, text in enumerate(texts, 1):
                    ws.append([f"段落{i}", text[:300]])  # 限制长度避免单元格溢出
                    ws.row_dimensions[i+6].height = 50
            
            # 图片内容
            if images:
                ws.append([])
                ws.append(["图片内容"])
                current_row = len(texts) + 8 if texts else 7
                for i, img_path in enumerate(images, 1):
                    try:
                        ws.cell(row=current_row, column=1, value=f"图片{i}")
                        img = ExcelImage(img_path)
                        img.width = 300
                        img.height = 200
                        ws.add_image(img, f"B{current_row}")
                        ws.row_dimensions[current_row].height = 180
                        current_row += 10  # 预留空间
                    except Exception as e:
                        print(f"图片 {i} 插入Excel失败：{str(e)}")
            
            excel_path = os.path.join(self.save_dir, "爬取数据汇总.xlsx")
            wb.save(excel_path)
            print(f"Excel表格已保存：{excel_path}")
        except Exception as e:
            print(f"保存Excel失败：{str(e)}")

    def run(self):
        """执行爬取流程"""
        try:
            if not self.load_page():
                return
            
            # 提取数据
            texts = self.extract_text()
            img_urls = self.extract_images()
            images = self.download_images(img_urls)
            
            # 保存数据
            if texts or images:
                self.save_to_word(texts, images)
                self.save_to_excel(texts, images)
                print("\n爬取完成！所有数据已保存")
            else:
                print("\n未提取到有效数据，请检查网页结构或调整配置参数")
                
        finally:
            if self.driver:
                self.driver.quit()
                print("浏览器已关闭")


if __name__ == "__main__":
    try:
        crawler = UniversalWebCrawler()
        crawler.run()
    except Exception as e:
        print(f"程序执行失败：{str(e)}")
